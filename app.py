from flask import Flask, request, jsonify, send_from_directory, send_file
from groq import Groq
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import io

app = Flask(__name__, static_folder='.')

GROQ_API_KEY = os.environ.get("GROQ_API_KEY")
client = Groq(api_key=GROQ_API_KEY)

@app.route('/')
def home():
    return send_from_directory('.', 'index.html')

@app.route('/logo1.gif')
def logo1():
    return send_from_directory('.', 'logo1.gif')

@app.route('/logo2.jpeg')
def logo2():
    return send_from_directory('.', 'logo2.jpeg')


FORMAT_GUIDES = {
    "Office Order": """
FORMAT: Office Order
- Opening: Start with "It is hereby ordered that..." OR "In exercise of the powers vested..."
- Use numbered paragraphs: 1., 2., 2.1, 2.2 etc.
- Language: Firm and directive. Third person throughout.
- Closing paragraph: "This issues with the approval of {authority}."
- Do NOT use "I" or personal pronouns.
- Example opening: "1. It is hereby ordered that all SSEs in-charge of sick line are directed to..."
""",

    "Office Memorandum": """
FORMAT: Office Memorandum (OM)
- Opening: Start with "The undersigned is directed to refer to..." OR "Attention is invited to..."
- Use numbered paragraphs: 1., 2., 2.1 etc.
- Language: Formal but communicative. Written as if from the office, not an individual.
- Closing paragraph: "This issues with the approval of {authority}."
- Do NOT use "I". Use "the undersigned" if referring to the writer.
- Example opening: "1. The undersigned is directed to refer to the subject cited above and to state that..."
""",

    "Circular": """
FORMAT: Circular
- Opening: Start with "Attention of all concerned is invited to..." OR "It has been observed that..." OR "Reference is invited to..."
- Use numbered paragraphs: 1., 2., 2.1 etc.
- Language: Informative and guidance-oriented. Addressed to a wide audience.
- Closing paragraph: "This circular issues with the approval of {authority}."
- May include sub-points as (a), (b), (c) under main paragraphs.
- Example opening: "1. Attention of all concerned is invited to the instructions contained in..."
""",

    "DO Letter": """
FORMAT: Demi Official (DO) Letter
- Opening: Start with "Dear [Designation]," on a new line, then begin the body.
- Do NOT use numbered paragraphs. Write in flowing prose paragraphs.
- Language: Semi-personal, formal but direct. Use "I" — written personally by the officer.
- Closing: End with "Yours sincerely," on a new line.
- Do NOT include "This issues with approval of..." — this is a personal letter.
- Keep it concise — typically 2-3 paragraphs only.
- Example opening: "Dear Sri/Smt [Name],\n\nI am writing to bring to your attention the matter of..."
""",

    "UO Note": """
FORMAT: Unofficial Note (UO Note)
- Opening: Start directly with the subject matter — no "To" address, no formal opening salutation.
- Do NOT use numbered paragraphs. Write as a brief internal note.
- Language: Brief, factual, and to the point. Used for internal departmental communication only.
- Closing: Simply end with the action requested or information conveyed. No formal closing line.
- Do NOT include "This issues with approval of..." — this is an internal note.
- Keep it very brief — typically 1-2 short paragraphs only.
- Example opening: "The matter regarding [subject] is brought to the notice of [department/officer]..."
"""
}


@app.route('/draft', methods=['POST'])
def draft():
    data         = request.json
    doc_type     = data.get('doc_type', 'Office Order')
    subject      = data.get('subject', '')
    references   = data.get('references', [])
    enclosures   = data.get('enclosures', [])
    instructions = data.get('instructions', '')
    addressees   = data.get('addressees', ['All Concerned'])
    authority    = data.get('for_officer', 'Sr. DME (Co)/BCT')
    tone         = data.get('tone', 'directive')

    tone_guide = {
        'directive':     'Use firm directive language.',
        'advisory':      'Use advisory, guidance-oriented language.',
        'clarificatory': 'This is a clarification of existing rules or instructions.',
        'reminder':      'This is a reminder/follow-up. Earlier instructions have not been complied with.'
    }

    addr_str = ', '.join(addressees) if isinstance(addressees, list) else addressees
    ref_str  = '; '.join(references) if references else 'None'
    custom_fmt = data.get('custom_format', '')
    fmt      = (custom_fmt if custom_fmt else FORMAT_GUIDES.get(doc_type, FORMAT_GUIDES["Office Order"])).format(authority=authority)

    prompt = f"""You are a senior Indian Railways officer in the C&W section of BCT Division, Western Railway, with 30 years of experience drafting official correspondence.

Your task is to write ONLY the body content of a {doc_type}.

{fmt}

STRICT RULES:
- Write ONLY the body content as described in the format above
- Do NOT include file number, date, subject line, reference lines at the top
- Do NOT include signature block or designation at the end (except "Yours sincerely," for DO Letter)
- Do NOT include Copy to section
- Do NOT include Enclosures section
- Do NOT write any heading or title before the body

Document details:
- Subject: {subject}
- References: {ref_str}
- Key instructions to convey: {instructions}
- Addressed to: {addr_str}
- Approving authority: {authority}
- Tone: {tone_guide.get(tone, tone_guide['directive'])}

Write the body content now:"""

    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=800
        )
        return jsonify({"text": response.choices[0].message.content})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/download', methods=['POST'])
def download():
    data       = request.json
    file_no    = data.get('file_no', '')
    date       = data.get('date', '')
    addressees = data.get('addressees', '')
    subject    = data.get('subject', '')
    references = data.get('references', [])
    enclosures = data.get('enclosures', [])
    body       = data.get('body', '')
    signed_by  = data.get('signed_by', 'ADME (C&W)/BCT')
    for_off    = data.get('for_officer', 'Sr. DME (Co)/BCT')
    copy_to    = data.get('copy_to', '')
    doc_type   = data.get('doc_type', 'Office Order')

    roman = ['i','ii','iii','iv','v','vi','vii','viii','ix','x']

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # Header table with logos
    header_table = doc.add_table(rows=1, cols=3)
    header_table.autofit = False
    header_table.columns[0].width = Inches(1.2)
    header_table.columns[1].width = Inches(4.1)
    header_table.columns[2].width = Inches(1.2)

    left_cell = header_table.cell(0, 0)
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        run = left_para.add_run()
        run.add_picture('logo1.gif', width=Inches(0.9))
    except:
        left_para.add_run('WR')

    center_cell = header_table.cell(0, 1)
    center_cell.paragraphs[0].clear()

    p_hindi = center_cell.add_paragraph('पश्चिम रेलवे')
    p_hindi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_hindi.runs[0]; r.font.size = Pt(16); r.font.bold = True

    p_eng = center_cell.add_paragraph('WESTERN RAILWAY')
    p_eng.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_eng.runs[0]; r.font.size = Pt(15); r.font.bold = True

    p_div = center_cell.add_paragraph('मंडल कार्यालय  |  Divisional Office')
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_div.runs[0].font.size = Pt(11)

    p_addr_lh = center_cell.add_paragraph('मुंबई सेंट्रल, मुंबई–400008  |  Mumbai Central, Mumbai-400008')
    p_addr_lh.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_addr_lh.runs[0].font.size = Pt(9)

    right_cell = header_table.cell(0, 2)
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        run = right_para.add_run()
        run.add_picture('logo2.jpeg', width=Inches(0.9))
    except:
        right_para.add_run('IR')

    doc.add_paragraph('')

    # File No and Date
    p_file = doc.add_paragraph()
    run_no = p_file.add_run(f'No. M/{file_no}')
    run_no.font.size = Pt(11); run_no.font.bold = True
    p_file.add_run('\t\t\t\t\t')
    run_date = p_file.add_run(f'Date: {date}')
    run_date.font.size = Pt(11)

    doc.add_paragraph('')

    # Addressed To — skip for UO Note
    if doc_type != 'UO Note' and addressees:
        addr_lines = addressees if isinstance(addressees, list) else [a.strip() for a in addressees.split('||') if a.strip()]
        for addr_line in addr_lines:
            p_to = doc.add_paragraph()
            p_to.add_run(addr_line).font.size = Pt(11)
        doc.add_paragraph('')

    # Subject
    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run(f'Subject: {subject}')
    run_sub.font.size = Pt(11); run_sub.font.bold = True

    # References
    if len(references) == 1:
        p_ref = doc.add_paragraph()
        p_ref.add_run(f'Reference: {references[0]}').font.size = Pt(11)
    elif len(references) > 1:
        for idx, ref in enumerate(references):
            p_ref = doc.add_paragraph()
            p_ref.add_run(f'Ref.({roman[idx] if idx < len(roman) else idx+1}): {ref}').font.size = Pt(11)

    doc.add_paragraph('')

    # Body paragraphs
    for line in body.split('\n'):
        if line.strip():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(line)
            run.font.size = Pt(11)
        else:
            doc.add_paragraph('')

    doc.add_paragraph('')
    doc.add_paragraph('')

    # Signature — DO Letter uses "Yours sincerely"
    if doc_type == 'DO Letter':
        p_sig0 = doc.add_paragraph()
        p_sig0.add_run('Yours sincerely,').font.size = Pt(11)
        doc.add_paragraph('')
        doc.add_paragraph('')

    p_sig1 = doc.add_paragraph()
    r1 = p_sig1.add_run(signed_by)
    r1.font.size = Pt(11); r1.font.bold = True

    # For officer line — not needed for DO Letter and UO Note
    if doc_type not in ['DO Letter', 'UO Note']:
        p_sig2 = doc.add_paragraph()
        r2 = p_sig2.add_run(f'For {for_off}')
        r2.font.size = Pt(11); r2.font.bold = True

    # Enclosures
    if enclosures:
        doc.add_paragraph('')
        for idx, encl in enumerate(enclosures):
            p_encl = doc.add_paragraph()
            label = f'Encl.({idx+1}): ' if len(enclosures) > 1 else 'Encl.: '
            p_encl.add_run(label + encl).font.size = Pt(11)

    doc.add_paragraph('')

    # Copy To — not for UO Note and DO Letter typically
    if copy_to and doc_type not in ['UO Note']:
        p_ct = doc.add_paragraph()
        r = p_ct.add_run('Copy to:')
        r.font.size = Pt(11); r.font.bold = True
        for line in copy_to.split('\n'):
            if line.strip():
                p_copy = doc.add_paragraph()
                p_copy.add_run(line.strip()).font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"BCT_CW_{subject[:30].replace(' ','_')}.docx"
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


if __name__ == '__main__':
    print("BCT Division C&W Circular Drafter is running!")
    print("Open your browser and go to: http://localhost:5000")
    app.run(debug=False, host='0.0.0.0', port=int(os.environ.get('PORT', 5000)))
